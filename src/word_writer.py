"""Word document writer based on python-docx."""

from __future__ import annotations

from pathlib import Path

from .config import DEFAULT_TITLE
from .utils import ensure_parent_dir


def write_table_to_docx(
    table_data: list[list[str]],
    output_path: str,
    title: str | None = None,
) -> None:
    """Write a 2D table matrix into an editable .docx table."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Please run: pip install -r requirements.txt"
        ) from exc

    path = ensure_parent_dir(output_path)
    document = Document()

    heading = document.add_heading(title or DEFAULT_TITLE, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not table_data:
        document.add_paragraph("未识别到有效表格内容")
        document.save(str(path))
        return

    max_cols = max(len(row) for row in table_data) if table_data else 1
    normalized_rows = [row + [""] * (max_cols - len(row)) for row in table_data]

    table = document.add_table(rows=len(normalized_rows), cols=max_cols)
    table.style = "Table Grid"

    for row_index, row in enumerate(normalized_rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value

    document.add_paragraph()
    document.add_paragraph("说明：带有 [低置信度] 标记的内容建议人工复核。")
    document.save(str(Path(path)))
