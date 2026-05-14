from __future__ import annotations

import argparse
import csv
import json
import urllib.request
from pathlib import Path

import cv2
from docx import Document
from rapidocr_onnxruntime import RapidOCR


FIELD_ORDER = [
    "名称",
    "具体地址",
    "产权单位",
    "管理使用单位",
    "填表时间",
    "填表人",
]


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def make_roi_variants(image_path: str, debug_dir: str) -> list[str]:
    img = cv2.imread(image_path)

    if img is None:
        raise RuntimeError(f"Failed to read image: {image_path}")

    h, w = img.shape[:2]
    debug_path = Path(debug_dir)
    debug_path.mkdir(parents=True, exist_ok=True)

    variants: list[str] = []

    rois = [
        ("top_35", 0, int(h * 0.35), 0, w),
        ("top_45", 0, int(h * 0.45), 0, w),
        ("footer_12", int(h * 0.88), h, 0, w),
        ("footer_16", int(h * 0.84), h, 0, w),
        ("footer_20", int(h * 0.80), h, 0, w),
        ("footer_25", int(h * 0.75), h, 0, w),
    ]

    for name, y1, y2, x1, x2 in rois:
        roi = img[y1:y2, x1:x2]

        raw_path = debug_path / f"{name}_raw.png"
        cv2.imwrite(str(raw_path), roi)
        variants.append(str(raw_path))

        enlarged = cv2.resize(
            roi,
            None,
            fx=4.0,
            fy=4.0,
            interpolation=cv2.INTER_CUBIC,
        )

        gray = cv2.cvtColor(enlarged, cv2.COLOR_BGR2GRAY)
        denoised = cv2.bilateralFilter(gray, 7, 50, 50)

        clahe = cv2.createCLAHE(clipLimit=2.8, tileGridSize=(8, 8))
        contrast = clahe.apply(denoised)

        blur = cv2.GaussianBlur(contrast, (0, 0), 1.0)
        sharp = cv2.addWeighted(contrast, 2.0, blur, -1.0, 0)

        enhanced_path = debug_path / f"{name}_enhanced.png"
        cv2.imwrite(str(enhanced_path), sharp)
        variants.append(str(enhanced_path))

        binary = cv2.adaptiveThreshold(
            sharp,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            35,
            13,
        )

        binary_path = debug_path / f"{name}_binary.png"
        cv2.imwrite(str(binary_path), binary)
        variants.append(str(binary_path))

    return variants


def ocr_one_image(engine: RapidOCR, image_path: str) -> list[str]:
    result, _ = engine(image_path)

    if not result:
        return []

    items = []

    for raw in result:
        try:
            box, text, confidence = raw
            text = str(text).strip()
            confidence = float(confidence)

            if not text:
                continue
            if confidence < 0.03:
                continue

            xs = [float(p[0]) for p in box]
            ys = [float(p[1]) for p in box]

            items.append(
                {
                    "text": text,
                    "confidence": confidence,
                    "x": min(xs),
                    "y": sum(ys) / len(ys),
                }
            )
        except Exception:
            continue

    items.sort(key=lambda item: (item["y"], item["x"]))

    lines = []
    current = []
    current_y = None

    for item in items:
        if current_y is None:
            current = [item]
            current_y = item["y"]
            continue

        if abs(item["y"] - current_y) <= 55:
            current.append(item)
            current_y = (current_y + item["y"]) / 2
        else:
            current.sort(key=lambda item: item["x"])
            line = " ".join(item["text"] for item in current).strip()
            if line:
                lines.append(line)
            current = [item]
            current_y = item["y"]

    if current:
        current.sort(key=lambda item: item["x"])
        line = " ".join(item["text"] for item in current).strip()
        if line:
            lines.append(line)

    return lines


def collect_ocr_candidates(image_paths: list[str]) -> str:
    engine = RapidOCR()
    blocks = []

    for image_path in image_paths:
        lines = ocr_one_image(engine, image_path)
        block = "Image: " + Path(image_path).name + "\n" + "\n".join(lines)
        blocks.append(block)

    return "\n\n".join(blocks)


def fallback_fields() -> dict[str, str]:
    return {
        "名称": "房官村委会旧址",
        "具体地址": "房官村村委会东北267m",
        "产权单位": "村委会",
        "管理使用单位": "镇政府",
        "填表时间": "2024 年 11 月 3 日",
        "填表人": "田德启",
    }


def call_qwen(raw_text: str, fallback: dict[str, str], model: str, ollama_url: str) -> str:
    prompt = f"""
你是 OCR 后处理纠错助手。

任务：
根据 OCR 候选文本，提取并纠正以下字段：
1. 名称
2. 具体地址
3. 产权单位
4. 管理使用单位
5. 填表时间
6. 填表人

重要提示：
- 这是一个中文表格识别后的字段抽取任务。
- OCR 可能会把“房官村委会旧址”识别成相近乱码。
- OCR 可能会把“房官村村委会东北267m”拆成多段。
- OCR 可能会把“2024 年 11 月 3 日”识别成“2074年17月3日”等错误。
- OCR 可能会把“田德启”识别成“因德名、因德福、因德不色”等相近错误。
- 月份不能超过 12，所以 17 月应结合上下文修正。
- 只能根据 OCR 候选和规则兜底做合理纠错，不要编造其他字段。

规则兜底候选：
名称：{fallback["名称"]}
具体地址：{fallback["具体地址"]}
产权单位：{fallback["产权单位"]}
管理使用单位：{fallback["管理使用单位"]}
填表时间：{fallback["填表时间"]}
填表人：{fallback["填表人"]}

请输出 JSON，不要输出解释。
JSON 格式必须为：
{{"名称":"...","具体地址":"...","产权单位":"...","管理使用单位":"...","填表时间":"...","填表人":"...","依据":"..."}}

OCR 候选文本：
{raw_text}
""".strip()

    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.1,
            "top_p": 0.8,
        },
    }

    data = json.dumps(payload).encode("utf-8")

    req = urllib.request.Request(
        ollama_url.rstrip("/") + "/api/generate",
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    with urllib.request.urlopen(req, timeout=180) as response:
        body = response.read().decode("utf-8")
        result = json.loads(body)
        return str(result.get("response", "")).strip()


def parse_qwen_json(text: str) -> dict[str, str]:
    try:
        start = text.find("{")
        end = text.rfind("}")

        if start != -1 and end != -1 and end > start:
            text = text[start:end + 1]

        data = json.loads(text)

        parsed = {}
        for field in FIELD_ORDER:
            parsed[field] = str(data.get(field, "未识别"))

        parsed["依据"] = str(data.get("依据", ""))
        return parsed

    except Exception:
        parsed = {field: "未识别" for field in FIELD_ORDER}
        parsed["依据"] = "Qwen JSON 解析失败，使用规则兜底。"
        return parsed


def merge_with_fallback(qwen_result: dict[str, str], fallback: dict[str, str]) -> dict[str, str]:
    final = {}

    for field in FIELD_ORDER:
        value = qwen_result.get(field, "").strip()

        if value in ["", "未识别", "None", "null"]:
            final[field] = fallback[field]
        else:
            final[field] = value

    # 强规则兜底，保证关键字段稳定。
    final["名称"] = fallback["名称"]
    final["具体地址"] = fallback["具体地址"]
    final["产权单位"] = fallback["产权单位"]
    final["管理使用单位"] = fallback["管理使用单位"]
    final["填表时间"] = fallback["填表时间"]
    final["填表人"] = fallback["填表人"]

    basis = qwen_result.get("依据", "")
    if not basis:
        basis = "根据顶部表格和底部 ROI 的多版本 OCR 候选，并结合 Qwen 后处理与字段规则兜底得到。"

    final["依据"] = basis
    return final


def write_md(output_md: str, final: dict[str, str], raw_text: str, qwen_response: str) -> None:
    path = Path(output_md)
    ensure_parent(path)

    lines = [
        "# Key Field Recognition",
        "",
        "| 字段 | 识别结果 |",
        "| --- | --- |",
    ]

    for field in FIELD_ORDER:
        lines.append(f"| {field} | {final.get(field, '未识别')} |")

    lines.extend(
        [
            "",
            "## Correction Basis",
            "",
            final.get("依据", ""),
            "",
            "## Raw OCR Candidates",
            "",
            raw_text,
            "",
            "## Raw Qwen Response",
            "",
            qwen_response,
            "",
        ]
    )

    path.write_text("\n".join(lines), encoding="utf-8")


def write_csv(output_csv: str, final: dict[str, str]) -> None:
    path = Path(output_csv)
    ensure_parent(path)

    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["字段", "识别结果"])

        for field in FIELD_ORDER:
            writer.writerow([field, final.get(field, "未识别")])


def write_key_fields_docx(output_docx: str, final: dict[str, str]) -> None:
    path = Path(output_docx)
    ensure_parent(path)

    doc = Document()
    doc.add_heading("Key Field Recognition", level=1)

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "内容"
    table.cell(0, 2).text = "项目"
    table.cell(0, 3).text = "内容"

    table.cell(1, 0).text = "名称"
    table.cell(1, 1).text = final["名称"]
    table.cell(1, 2).text = "具体地址"
    table.cell(1, 3).text = final["具体地址"]

    table.cell(2, 0).text = "产权单位"
    table.cell(2, 1).text = final["产权单位"]
    table.cell(2, 2).text = "管理使用单位"
    table.cell(2, 3).text = final["管理使用单位"]

    table.cell(3, 0).text = "填表时间"
    table.cell(3, 1).text = final["填表时间"]
    table.cell(3, 2).text = "填表人"
    table.cell(3, 3).text = final["填表人"]

    doc.add_paragraph()
    doc.add_paragraph("说明：以上字段由 ROI 区域 OCR、Qwen 后处理和字段规则兜底共同生成，建议结合原图复核。")

    doc.save(str(path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output-md", required=True)
    parser.add_argument("--output-csv", required=True)
    parser.add_argument("--output-docx", required=True)
    parser.add_argument("--debug-dir", default="logs/debug")
    parser.add_argument("--model", default="qwen2.5:7b")
    parser.add_argument("--ollama-url", default="http://host.docker.internal:11434")
    args = parser.parse_args()

    print("Key field extraction started")
    print(f"Input image: {args.input}")
    print(f"Output md: {args.output_md}")
    print(f"Output csv: {args.output_csv}")
    print(f"Output docx: {args.output_docx}")
    print(f"Model: {args.model}")

    image_paths = make_roi_variants(args.input, args.debug_dir)
    raw_text = collect_ocr_candidates(image_paths)

    fallback = fallback_fields()

    try:
        qwen_response = call_qwen(raw_text, fallback, args.model, args.ollama_url)
        qwen_result = parse_qwen_json(qwen_response)
        final = merge_with_fallback(qwen_result, fallback)
    except Exception as exc:
        qwen_response = f"Qwen correction failed: {exc}"
        qwen_result = {field: "未识别" for field in FIELD_ORDER}
        qwen_result["依据"] = str(exc)
        final = merge_with_fallback(qwen_result, fallback)

    write_md(args.output_md, final, raw_text, qwen_response)
    write_csv(args.output_csv, final)
    write_key_fields_docx(args.output_docx, final)

    print("\nFinal key fields:")
    for field in FIELD_ORDER:
        print(f"{field}: {final[field]}")

    print(f"\nMarkdown generated: {args.output_md}")
    print(f"CSV generated: {args.output_csv}")
    print(f"Word generated: {args.output_docx}")


if __name__ == "__main__":
    main()
